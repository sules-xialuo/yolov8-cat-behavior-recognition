from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "1-参考文献与图表题注修正版.docx"
OUTPUT = ROOT / "1-引用位置与参考文献格式修正版.docx"


def set_run_font(run, size=10.5):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def set_text(paragraph, text, size=10.5):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_snap_to_grid(paragraph, enabled=False):
    p_pr = paragraph._p.get_or_add_pPr()
    snap = p_pr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        p_pr.append(snap)
    snap.set(qn("w:val"), "1" if enabled else "0")


def format_reference_entry(paragraph):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_snap_to_grid(paragraph, False)
    for run in paragraph.runs:
        set_run_font(run, 10.5)


def format_heading(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_snap_to_grid(paragraph, False)


def main():
    doc = Document(SOURCE)

    updates = {
        80: "在现代，宠物在家庭生活中的陪伴价值不断提高，但猫类情绪和行为状态通常无法较好的通过语言直接表达，饲养人主要依据耳朵方向、瞳孔形态、尾巴姿态和身体动作等外显特征进行判断。这种判断方式依赖个人经验，在复杂环境、光照、侧面姿态和多猫场景中容易产生偏差。",
        81: "深度学习目标检测算法能够从图像中自动定位并分类目标，为宠物行为识别提供了新的技术路径。与识别整只猫的类别不同，本文关注与行为判断直接相关的局部状态，例如耳朵前倾、圆瞳、竖瞳、尾巴上抬、尾巴下垂和尾巴蜷缩等。通过对这些局部特征进行检测，可以进一步组合推断警觉、放松、不安等状态。",
        88: "本文研究内容包括：构建七类猫局部状态数据集；设计离线数据增强与重点样本重采样策略；基于YOLOv8s进行迁移学习和多轮微调；设计疑似标签不一致样本发现、LabelImg复核和标签同步流程；实现桌面端和网页端单图检测界面；最后根据训练结果与应用部署效果分析模型效果和不足。",
        93: "目标检测需要同时预测目标类别和位置。本文采用YOLO格式标注，每张图像对应一个同名txt文件，每行由类别编号、中心点横坐标、中心点纵坐标、宽度和高度组成，坐标均按图像宽高归一化。该格式结构简单，以便与Ultralytics YOLO训练框架配合[1-3]。",
        95: "YOLOv8是Ultralytics推出的一阶段目标检测框架，具有训练接口统一、推理速度较快、部署生态完善等特点。本文选用YOLOv8s作为主要模型，在精度和显存占用之间取得折中。模型通过主干网络提取多尺度特征，再由检测头输出边界框、类别和置信度[3-5]。",
        97: "自建猫状态数据集存在样本数量有限、类别不均衡和目标尺度较小等问题。本文使用水平翻转、仿射变换、透视变换、亮度对比度扰动、色调饱和度变化、CLAHE、轻微模糊和噪声等增强方式。同时利用预训练权重进行迁移学习，先训练增强基线模型，再针对困难样本和修正后的标签进行小学习率微调[6-8]。",
        135: "图形界面由jiance.py实现，使用Tkinter提供图片选择、结果文本显示和清空功能。用户选择图片后，程序加载YOLO模型进行预测，使用result.plot()绘制检测框，并通过OpenCV和Pillow完成图像格式转换与界面显示。文本区域输出目标数量、类别名称、置信度和边界框坐标[5,9]。",
        143: "实验在Windows环境下完成，训练框架为Ultralytics YOLO，图像处理使用OpenCV、Pillow和Albumentations。训练使用单卡GPU，batch size设置为8，workers设置为0以适配Windows本地训练环境。所有训练均保存results.csv、训练曲线、PR曲线、F1曲线、混淆矩阵和验证集预测样例[5,8-9]。",
        167: "从疑似错标统计和混淆矩阵可以看出，耳朵前倾与耳朵双边之间仍是主要混淆来源，圆瞳与竖瞳也存在着一定类别冲突。这些问题一方面来自猫局部状态本身边界模糊，另一方面来自图像分辨率、遮挡和标注标准不完全一致。尾巴类状态受姿态和可见区域影响明显，部分情况下的尾巴下垂、尾巴蜷缩和其他身体部位边界较难区分[10-13]。",
        173: "实验结果显示，最新模型cats_focus_yolov8s6取得mAP50=0.8296、mAP50-95=0.6472，明显优于增强基线模型。该结果说明，在猫局部状态识别这种细粒度任务中，预训练迁移、类别重采样、重点样本微调和标签质量控制具有实际效果。",
        177: "（3）行为推理仍较简单：当前规则主要基于单张图片，未来可加入视频时序、姿态估计和概率推理[13]。",
    }
    for idx, text in updates.items():
        set_text(doc.paragraphs[idx], text)

    refs = [
        "[1] 赵永强, 饶元, 董世鹏, 等. 深度学习目标检测方法综述[J]. 中国图象图形学报, 2020, 25(4): 629-654.",
        "[2] 曹家乐, 李亚利, 孙汉卿, 等. 基于深度学习的视觉目标检测技术综述[J]. 中国图象图形学报, 2022, 27(6): 1697-1722.",
        "[3] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788.",
        "[4] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal Speed and Accuracy of Object Detection[EB/OL]. arXiv:2004.10934, 2020.",
        "[5] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLO[CP/OL]. https://github.com/ultralytics/ultralytics, 2023.",
        "[6] 冯晓硕, 沈樾, 王冬琦. 基于图像的数据增强方法发展现状综述[J]. 计算机科学与应用, 2021, 11(2): 370-382.",
        "[7] 郭文明, 郭鸣, 王俊, 等. 类激活映射指导数据增强的细粒度图像分类[J]. 计算机辅助设计与图形学学报, 2021, 33(11): 1698-1704.",
        "[8] Buslaev A, Iglovikov V I, Khvedchenya E, et al. Albumentations: Fast and Flexible Image Augmentations[J]. Information, 2020, 11(2): 125.",
        "[9] Bradski G. The OpenCV Library[J]. Dr. Dobb's Journal of Software Tools, 2000, 25(11): 120-123.",
        "[10] 李林葳, 宋鑫悦, 张智盛, 等. 基于图像处理技术的动物行为识别研究进展[J]. 中国畜牧杂志, 2024, 60(10): 24-34.",
        "[11] 朱芷芫, 王海峰, 李斌, 等. 深度学习在畜禽典型行为识别中的研究进展[J]. 中国农业科技导报, 2024, 26(10): 110-124.",
        "[12] Evangelista M C, Watanabe R, Leung V S Y, et al. Facial expressions of pain in cats: the development and validation of a Feline Grimace Scale[J]. Scientific Reports, 2019, 9: 19128.",
        "[13] Pereira T D, Tabris N, Matsliah A, et al. SLEAP: A deep learning system for multi-animal pose tracking[J]. Nature Methods, 2022, 19(4): 486-495.",
    ]

    thanks_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "致谢")
    ref_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")
    thanks_text = doc.paragraphs[thanks_idx + 1].text

    start = min(thanks_idx, ref_idx)
    for p in list(doc.paragraphs[start:]):
        remove_paragraph(p)

    anchor = doc.paragraphs[-1]
    ref_heading = insert_after(anchor, "参考文献", style="Heading 1")
    format_heading(ref_heading)
    anchor = ref_heading
    for ref in refs:
        anchor = insert_after(anchor, ref, style="Normal")
        format_reference_entry(anchor)
    anchor = insert_after(anchor, "")
    thanks_heading = insert_after(anchor, "致谢", style="Heading 1")
    format_heading(thanks_heading)
    thanks_para = insert_after(thanks_heading, thanks_text, style="Normal")
    format_reference_entry(thanks_para)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
