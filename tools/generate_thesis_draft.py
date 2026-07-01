from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "陈祖烨-基于YOLOv8的猫行为识别系统论文初稿-重构版.docx"
RUNS_DIR = ROOT / "runs" / "detect"
BEST_RUN = RUNS_DIR / "cats_focus_yolov8s6"
BASE_RUN = RUNS_DIR / "cats_aug_yolov8s2"
OLD_FOCUS_RUN = RUNS_DIR / "cats_focus_yolov8s"
WEB_OUTPUT_DIR = ROOT / "runs" / "cat_web_app"

NAMES = ["shangtai", "qianqing", "xiachui", "shuangbian", "yuantong", "quansuo", "shutong"]
CN_NAMES = ["尾巴上抬", "耳朵前倾", "尾巴下垂", "耳朵双边", "圆瞳", "尾巴蜷缩", "竖瞳"]


def font_run(run, font: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_para_format(p, first_line: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(24) if first_line else Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_toc(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def setup_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.75)
    add_page_number(sec.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    heading_specs = [
        ("Heading 1", "黑体", 16, True, 12, 6),
        ("Heading 2", "黑体", 14, True, 6, 3),
        ("Heading 3", "黑体", 12, True, 3, 0),
    ]
    for name, font, size, bold, before, after in heading_specs:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def center(doc: Document, text: str, size: float, font: str = "宋体", bold: bool = False):
    p = doc.add_paragraph()
    set_para_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    font_run(run, font, size, bold)
    return p


def body(doc: Document, text: str):
    p = doc.add_paragraph()
    set_para_format(p)
    run = p.add_run(text)
    font_run(run, "宋体", 12)
    return p


def noindent(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    set_para_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    font_run(run, "宋体", 12, bold)
    return p


def item(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    set_para_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent = Pt(24)
    r1 = p.add_run(label)
    font_run(r1, "宋体", 12, True)
    r2 = p.add_run(text)
    font_run(r2, "宋体", 12)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    font_run(run, "宋体", 10.5, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == len(row) - 1 and len(row) > 3 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(t.cell(i, j), text, bold=i == 0, align=align)
            if i == 0:
                shade_cell(t.cell(i, j), "EDEDED")
            if widths:
                t.cell(i, j).width = Cm(widths[j])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Pt(0)


def figure(doc: Document, path: Path, caption: str, width_cm: float = 13.8) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    set_para_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = center(doc, caption, 10.5, "宋体", False)
    cap.paragraph_format.space_after = Pt(3)


def read_best(run_dir: Path) -> dict[str, str]:
    rows = list(csv.DictReader((run_dir / "results.csv").open(encoding="utf-8")))
    return max(rows, key=lambda r: float(r["metrics/mAP50-95(B)"]))


def read_last(run_dir: Path) -> dict[str, str]:
    rows = list(csv.DictReader((run_dir / "results.csv").open(encoding="utf-8")))
    return rows[-1]


def count_dataset(ds: Path) -> dict[str, dict[str, object]]:
    out: dict[str, dict[str, object]] = {}
    for split in ("train", "val"):
        img_dir = ds / "images" / split
        lab_dir = ds / "labels" / split
        counts = [0] * len(NAMES)
        for label in lab_dir.glob("*.txt"):
            for line in label.read_text(encoding="utf-8", errors="ignore").splitlines():
                parts = line.split()
                if len(parts) == 5 and parts[0].isdigit():
                    cls = int(parts[0])
                    if 0 <= cls < len(counts):
                        counts[cls] += 1
        out[split] = {
            "images": sum(1 for p in img_dir.iterdir() if p.is_file()) if img_dir.exists() else 0,
            "labels": sum(1 for p in lab_dir.glob("*.txt")) if lab_dir.exists() else 0,
            "boxes": sum(counts),
            "counts": counts,
        }
    return out


def run_summary_rows() -> list[list[str]]:
    rows = [["训练阶段", "run目录", "最佳epoch", "Precision", "Recall", "mAP50", "mAP50-95"]]
    for run in sorted(RUNS_DIR.glob("cats_*")):
        if not (run / "results.csv").exists():
            continue
        best = read_best(run)
        rows.append(
            [
                "增强基线" if "aug" in run.name else "重点微调",
                run.name,
                best["epoch"],
                f"{float(best['metrics/precision(B)']):.4f}",
                f"{float(best['metrics/recall(B)']):.4f}",
                f"{float(best['metrics/mAP50(B)']):.4f}",
                f"{float(best['metrics/mAP50-95(B)']):.4f}",
            ]
        )
    return rows


def suspect_pair_rows(path: Path) -> list[list[str]]:
    if not path.exists():
        return []
    rows = list(csv.DictReader(path.open(encoding="utf-8-sig")))
    pairs: dict[tuple[str, str], int] = {}
    for row in rows:
        key = (row["gt_class"], row["pred_class"])
        pairs[key] = pairs.get(key, 0) + 1
    out = [["人工标签类别", "模型预测类别", "疑似数量"]]
    for (gt, pred), n in sorted(pairs.items(), key=lambda item_: item_[1], reverse=True)[:8]:
        out.append([gt, pred, str(n)])
    return out


def latest_web_output() -> Path | None:
    if not WEB_OUTPUT_DIR.exists():
        return None
    images = sorted(
        [p for p in WEB_OUTPUT_DIR.glob("*.jpg") if p.is_file()],
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return images[0] if images else None


def build() -> None:
    doc = Document()
    setup_document(doc)

    cats = count_dataset(ROOT / "datasets" / "cats")
    aug = count_dataset(ROOT / "datasets" / "cats_aug")
    focus = count_dataset(ROOT / "datasets" / "cats_focus")
    base_best = read_best(BASE_RUN)
    old_best = read_best(OLD_FOCUS_RUN)
    best = read_best(BEST_RUN)
    last = read_last(BEST_RUN)
    suspect_csv = BEST_RUN / "suspect_label_mismatches.csv"
    suspect_count = sum(1 for _ in csv.DictReader(suspect_csv.open(encoding="utf-8-sig"))) if suspect_csv.exists() else 0
    web_outputs = sum(1 for p in WEB_OUTPUT_DIR.glob("*.jpg") if p.is_file()) if WEB_OUTPUT_DIR.exists() else 0
    web_latest = latest_web_output()

    center(doc, "毕业设计（论文）", 22, "黑体", True)
    for _ in range(3):
        doc.add_paragraph()
    center(doc, "基于YOLOv8的猫行为识别系统设计与实现", 18, "黑体", True)
    for _ in range(5):
        doc.add_paragraph()
    table(
        doc,
        [
            ["学生姓名", "陈祖烨"],
            ["学号", "23687"],
            ["专业班级", "待补充"],
            ["指导教师", "待补充"],
            ["学院", "待补充"],
            ["完成日期", "2026年5月"],
        ],
        widths=[4.0, 8.0],
    )
    doc.add_page_break()

    center(doc, "摘  要", 16, "黑体", True)
    body(
        doc,
        "随着计算机视觉和深度学习技术的发展，目标检测方法逐渐应用于智能监控、动物行为分析和宠物健康管理等场景。"
        "针对宠物猫行为理解主要依赖人工经验、判断标准不统一等问题，本文设计并实现了一套基于YOLOv8的猫行为识别系统。"
        "系统以猫耳、猫眼和猫尾等关键局部特征为检测对象，将尾巴上抬、耳朵前倾、尾巴下垂、耳朵双边、圆瞳、尾巴蜷缩和竖瞳七种状态作为识别类别，"
        "通过目标检测结果为后续行为状态推理提供结构化输入。"
    )
    body(
        doc,
        f"本文在原始自建数据集基础上进行了多轮数据扩充、标签复核和困难样本挖掘。更新后的原始数据集包含训练图像{cats['train']['images']}张、验证图像{cats['val']['images']}张，"
        f"重点微调数据集包含训练图像{focus['train']['images']}张、目标框{focus['train']['boxes']}个。"
        "在训练流程上，本文先构建离线增强数据集训练基线模型，再利用人工挑选样本、类别重采样和标签纠错后的数据进行多轮微调。"
        f"最新微调实验cats_focus_yolov8s6在第{best['epoch']}轮取得最佳结果，Precision为{float(best['metrics/precision(B)']):.4f}，"
        f"Recall为{float(best['metrics/recall(B)']):.4f}，mAP50为{float(best['metrics/mAP50(B)']):.4f}，"
        f"mAP50-95为{float(best['metrics/mAP50-95(B)']):.4f}。"
        f"相较增强基线模型mAP50-95={float(base_best['metrics/mAP50-95(B)']):.4f}，检测精度显著提升。"
    )
    body(
        doc,
        "此外，本文实现了基于模型预测与人工标签高IoU冲突的疑似错标样本发现流程，并结合LabelImg复核、标签同步和备份机制完成闭环数据清洗。"
        "在应用层面，系统新增了基于本地网页服务的部署方式，支持浏览器上传图片、返回检测框、输出局部状态和生成照护建议。"
        "实验表明，对于小目标、局部遮挡和细粒度类别易混淆的猫状态识别任务，数据质量控制、二阶段微调和可交互应用封装比单纯增加训练轮次更关键。"
    )
    noindent(doc, "关键词：YOLOv8；目标检测；猫行为识别；数据增强；标签清洗；迁移学习")

    doc.add_page_break()
    center(doc, "Abstract", 16, "Times New Roman", True)
    body(
        doc,
        "With the development of computer vision and deep learning, object detection has become an effective tool for animal behavior analysis. "
        "This thesis designs and implements a cat behavior recognition system based on YOLOv8. The system detects seven fine-grained local states of cat ears, eyes and tail, "
        "including raised tail, forward ears, lowered tail, bilateral ears, round pupil, curled tail and vertical pupil. "
        "These detection results provide structured visual evidence for rule-based behavior inference."
    )
    body(
        doc,
        f"The updated dataset contains {cats['train']['images']} training images and {cats['val']['images']} validation images in the original set. "
        f"After offline augmentation, focus-sample resampling and label review, the final fine-tuning set contains {focus['train']['images']} training images and {focus['train']['boxes']} boxes. "
        f"The latest fine-tuned run achieves {float(best['metrics/mAP50(B)']):.4f} mAP50 and {float(best['metrics/mAP50-95(B)']):.4f} mAP50-95, "
        "showing that label cleaning and low-learning-rate fine-tuning can improve fine-grained cat-state detection."
    )
    noindent(doc, "Key words: YOLOv8; object detection; cat behavior recognition; data augmentation; label cleaning; transfer learning")

    doc.add_page_break()
    center(doc, "目  录", 16, "黑体", True)
    add_toc(doc.add_paragraph())
    noindent(doc, "提示：在 Microsoft Word 中右键目录并选择“更新域”，即可生成正式页码。")
    doc.add_page_break()

    doc.add_heading("第一章 绪论", level=1)
    doc.add_heading("1.1 研究背景", level=2)
    body(
        doc,
        "宠物猫在家庭生活中的陪伴价值不断提高，但猫的情绪和行为状态通常无法通过语言直接表达，饲养者主要依据耳朵方向、瞳孔形态、尾巴姿态和身体动作等外显特征进行判断。"
        "这种判断方式依赖个人经验，在复杂光照、遮挡、侧面姿态和多猫场景中容易产生偏差。"
    )
    body(
        doc,
        "深度学习目标检测算法能够从图像中自动定位并分类目标，为宠物行为识别提供了新的技术路径。与识别整只猫的类别不同，本文关注与行为判断直接相关的局部状态，"
        "例如耳朵前倾、圆瞳、竖瞳、尾巴上抬、尾巴下垂和尾巴蜷缩等。通过对这些局部特征进行检测，可以进一步组合推断警觉、放松、不安等状态。"
    )
    doc.add_heading("1.2 研究目的与意义", level=2)
    body(
        doc,
        "本文目标是构建一套从数据采集、标注、增强、训练、标签清洗到图像推理展示的猫行为识别系统。"
        "研究重点不是直接给出完全可靠的动物情绪结论，而是先建立可复现的局部状态识别流程，为后续行为推理和多模态扩展打下基础。"
    )
    item(doc, "（1）理论意义：", "探索目标检测技术与宠物局部行为学特征结合的方法，为细粒度动物行为识别提供工程案例。")
    item(doc, "（2）实践意义：", "通过自动检测猫耳、猫眼和猫尾状态，辅助用户理解宠物状态，提高人与宠物互动质量。")
    item(doc, "（3）应用意义：", "该流程可迁移到犬类、流浪动物监测和智能养殖场景，具有一定推广价值。")
    doc.add_heading("1.3 研究内容", level=2)
    body(
        doc,
        "本文研究内容包括：构建七类猫局部状态数据集；设计离线数据增强与重点样本重采样策略；基于YOLOv8s进行迁移学习和多轮微调；"
        "设计疑似标签不一致样本发现、LabelImg复核和标签同步流程；实现桌面端和网页端单图检测界面；最后根据训练结果与应用部署效果分析模型效果和不足。"
    )
    doc.add_heading("1.4 论文结构", level=2)
    body(doc, "论文共六章：第一章介绍研究背景和目标；第二章介绍相关技术；第三章给出系统需求和总体设计；第四章描述数据处理和模型训练实现；第五章分析实验结果；第六章总结与展望。")

    doc.add_heading("第二章 相关技术基础", level=1)
    doc.add_heading("2.1 目标检测与YOLO格式标注", level=2)
    body(
        doc,
        "目标检测需要同时预测目标类别和位置。本文采用YOLO格式标注，每张图像对应一个同名txt文件，每行由类别编号、中心点横坐标、中心点纵坐标、宽度和高度组成，"
        "坐标均按图像宽高归一化。该格式结构简单，便于与Ultralytics YOLO训练框架配合。"
    )
    doc.add_heading("2.2 YOLOv8模型", level=2)
    body(
        doc,
        "YOLOv8是Ultralytics推出的一阶段目标检测框架，具有训练接口统一、推理速度较快、部署生态完善等特点。"
        "本文选用YOLOv8s作为主要模型，在精度和显存占用之间取得折中。模型通过主干网络提取多尺度特征，再由检测头输出边界框、类别和置信度。"
    )
    doc.add_heading("2.3 数据增强与迁移学习", level=2)
    body(
        doc,
        "自建猫状态数据集存在样本数量有限、类别不均衡和目标尺度较小等问题。本文使用水平翻转、仿射变换、透视变换、亮度对比度扰动、色调饱和度变化、CLAHE、轻微模糊和噪声等增强方式。"
        "同时利用预训练权重进行迁移学习，先训练增强基线模型，再针对困难样本和修正后的标签进行小学习率微调。"
    )
    doc.add_heading("2.4 评价指标", level=2)
    body(
        doc,
        "实验采用Precision、Recall、mAP50和mAP50-95评价模型。Precision反映预测结果中正确目标比例，Recall反映真实目标被检出的比例。"
        "mAP50使用IoU阈值0.5计算平均精度，mAP50-95在多个IoU阈值上取平均，对定位质量要求更高。"
    )

    doc.add_heading("第三章 系统需求分析与总体设计", level=1)
    doc.add_heading("3.1 功能需求", level=2)
    body(
        doc,
        "系统需要支持数据集构建、数据增强、模型训练、训练结果汇总、疑似错标样本发现、人工复核同步、单图检测和结果展示等功能。"
        "其中，数据质量控制和网页应用部署是本次更新后的重点功能：前者用于减少细粒度类别错标对训练造成的干扰，后者用于把模型能力封装为更易使用的浏览器交互系统。"
    )
    table(
        doc,
        [
            ["模块", "功能", "对应文件"],
            ["数据合并", "将新增原始标注图像按哈希方式合并到训练/验证集", "tools/merge_original_dataset.py"],
            ["离线增强", "对中低频类别做增强复制，提升数据多样性", "tools/augment_cats_dataset.py"],
            ["重点集构建", "加入人工挑选样本，并对弱势类别重复采样", "tools/build_cats_focus_dataset.py"],
            ["模型训练", "加载权重并进行YOLOv8s微调", "train_cats_focus.py"],
            ["错标发现", "通过高IoU但类别冲突的预测发现疑似标签问题", "tools/find_suspect_label_mismatches.py"],
            ["复核同步", "准备LabelImg复核目录，并将修正标签同步回数据集", "tools/prepare_labelimg_suspect_review.py / sync_reviewed_labels.py"],
            ["应用检测", "选择图片并展示类别、置信度和检测框", "jiance.py"],
            ["网页部署", "提供浏览器上传、模型推理、结果图片返回、人工状态输入和照护建议生成", "cat_web_app.py"],
        ],
        widths=[2.8, 6.3, 5.2],
    )
    doc.add_heading("3.2 系统流程设计", level=2)
    body(
        doc,
        "系统整体流程为：首先收集和标注猫图像，使用合并脚本扩充原始数据集；随后进行数据增强和重点样本构建；"
        "然后训练增强基线模型和重点微调模型；训练后使用模型扫描训练/验证图像，找出高IoU但类别不一致的样本；"
        "将疑似样本导入LabelImg进行人工修正，再同步回原始数据集；最后重新训练，并通过桌面界面或网页服务完成单图推理和照护建议输出。"
    )
    table(
        doc,
        [
            ["阶段", "输入", "处理", "输出"],
            ["数据扩充", "原始图片与YOLO标签", "合法性校验、哈希划分、复制", "datasets/cats"],
            ["数据增强", "datasets/cats", "类别重采样和图像扰动", "datasets/cats_aug"],
            ["重点微调集", "增强集和精选样本", "弱势类重复采样", "datasets/cats_focus"],
            ["训练评估", "数据集和预训练权重", "YOLOv8s训练、验证和绘图", "best.pt、results.csv、曲线图"],
            ["标签清洗", "模型预测与人工标签", "高IoU类别冲突检测、人工复核", "修正后的标签和备份"],
            ["应用推理", "用户图片和最终权重", "检测框绘制、状态组合和建议生成", "桌面界面与网页端结果"],
        ],
        widths=[2.5, 3.5, 5.0, 3.5],
    )
    doc.add_heading("3.3 行为推理规则设计", level=2)
    body(
        doc,
        "本文当前系统主要完成局部状态检测，行为推理采用规则化设计。检测结果可按耳朵、眼睛和尾巴进行组合，当多个局部状态置信度超过阈值时输出候选行为状态。"
        "由于单张图片无法完全表达连续行为，规则推理结果应作为辅助提示而非绝对结论。"
    )
    table(
        doc,
        [
            ["局部状态组合", "候选状态", "说明"],
            ["耳朵前倾 + 竖瞳 + 尾巴下垂", "警觉或紧张", "通常表示注意力集中或受到外界刺激"],
            ["耳朵双边 + 圆瞳 + 尾巴上抬", "放松或亲近", "状态较平稳，可能具备互动倾向"],
            ["尾巴蜷缩 + 竖瞳", "不安或防御", "可能处于压力、回避或防御状态"],
            ["圆瞳 + 尾巴上抬", "好奇或探索", "需要结合耳朵方向和场景进一步判断"],
        ],
        widths=[5.8, 3.2, 5.0],
    )

    doc.add_heading("第四章 系统实现", level=1)
    doc.add_heading("4.1 数据集更新与类别分布", level=2)
    body(
        doc,
        f"本次项目更新后，原始数据集训练集包含{cats['train']['images']}张图像、{cats['train']['boxes']}个目标框，验证集包含{cats['val']['images']}张图像、{cats['val']['boxes']}个目标框。"
        f"离线增强后训练集扩展为{aug['train']['images']}张图像、{aug['train']['boxes']}个目标框；重点微调集训练部分进一步扩展为{focus['train']['images']}张图像、{focus['train']['boxes']}个目标框。"
        "验证集在不同阶段保持一致，便于横向比较模型效果。"
    )
    table(
        doc,
        [["类别编号", "英文标签", "中文含义", "原始训练框", "验证框", "增强训练框", "重点训练框"]]
        + [
            [
                str(i),
                NAMES[i],
                CN_NAMES[i],
                str(cats["train"]["counts"][i]),
                str(cats["val"]["counts"][i]),
                str(aug["train"]["counts"][i]),
                str(focus["train"]["counts"][i]),
            ]
            for i in range(len(NAMES))
        ],
        widths=[1.4, 2.5, 2.3, 2.0, 1.8, 2.2, 2.2],
    )
    figure(doc, BEST_RUN / "labels.jpg", "图4-1 最新重点数据集标签分布与标注框示例", 13.2)

    doc.add_heading("4.2 标签清洗流程实现", level=2)
    body(
        doc,
        "细粒度猫状态识别对标注质量较敏感。例如耳朵前倾与耳朵双边、圆瞳与竖瞳在局部遮挡或低分辨率条件下容易被混淆。"
        "为降低错标影响，本文新增了疑似标签不一致样本发现流程。该流程使用已训练模型扫描图像，若某个人工标注框与模型预测框IoU较高，但类别不同且预测置信度较高，则将其记录为疑似错标样本。"
        f"最新run中共导出{suspect_count}条疑似标签冲突记录。"
    )
    pair_rows = suspect_pair_rows(suspect_csv)
    if pair_rows:
        table(doc, pair_rows, widths=[4.0, 4.0, 2.5])
    body(
        doc,
        "疑似样本随后通过prepare_labelimg_suspect_review.py整理为LabelImg可直接打开的复核目录，并生成manifest.csv记录原图、原标签、复核图和复核标签之间的映射关系。"
        "人工修改后，sync_reviewed_labels.py会比较复核标签与原始标签差异，将变化同步回源数据集，同时在review目录下创建时间戳备份，避免误覆盖。"
    )
    table(
        doc,
        [
            ["复核批次", "manifest记录数", "说明"],
            ["suspect_labels", "67", "第一轮高置信疑似错标复核"],
            ["suspect_labels_remaining", "140", "第一轮后剩余疑似样本整理"],
            ["suspect_labels_round2", "47", "第二轮疑似样本复核"],
        ],
        widths=[4.0, 3.0, 6.0],
    )
    doc.add_heading("4.3 模型训练实现", level=2)
    body(
        doc,
        "最新训练脚本train_cats_focus.py加载weights/cats_focus_best.pt，在yolo-cats-focus.yaml指定的数据集上进行小学习率微调。"
        "主要参数为epochs=12，patience=4，imgsz=768，batch=8，optimizer=AdamW，lr0=0.00005，lrf=0.01，cos_lr=True，warmup_epochs=0，"
        "weight_decay=0.0005，box=7.5，cls=0.8，dfl=1.5，mosaic=0，mixup=0，degrees=1.0，translate=0.02，scale=0.1，fliplr=0.5。"
        "与早期微调相比，本轮训练更强调在已清洗标签基础上的稳定微调，减少过强增强对小目标定位造成的扰动。"
    )
    table(
        doc,
        [
            ["参数", "增强基线训练", "早期重点微调", "最新重点微调"],
            ["权重来源", "YOLOv8s/续训权重", "cats_aug_baseline_best.pt", "cats_focus_best.pt"],
            ["数据集", "yolo-cats-aug.yaml", "yolo-cats-focus.yaml", "yolo-cats-focus.yaml"],
            ["epochs / patience", "100 / 25", "40 / 12", "12 / 4"],
            ["optimizer", "auto", "auto", "AdamW"],
            ["lr0", "0.002", "0.0005", "0.00005"],
            ["imgsz / batch", "768 / 8", "768 / 8", "768 / 8"],
            ["mosaic / mixup", "0.1 / 0", "0 / 0", "0 / 0"],
        ],
        widths=[3.2, 3.8, 3.8, 3.8],
    )
    doc.add_heading("4.4 应用界面与网页部署实现", level=2)
    body(
        doc,
        "图形界面由jiance.py实现，使用Tkinter提供图片选择、结果文本显示和清空功能。用户选择图片后，程序加载YOLO模型进行预测，"
        "使用result.plot()绘制检测框，并通过OpenCV和Pillow完成图像格式转换与界面显示。文本区域输出目标数量、类别名称、置信度和边界框坐标。"
    )
    body(
        doc,
        "为提升系统可用性，本次更新新增cat_web_app.py，将检测模型部署为本地网页应用。网页服务基于Python标准库ThreadingHTTPServer实现，默认监听127.0.0.1:7860。"
        "前端页面提供图片上传、识别按钮、结果预览、检测目标列表、人工观察输入和照护建议展示等区域；后端提供/api/predict和/api/manual两个接口。"
        "/api/predict接收multipart图片文件，调用cats_focus_yolov8s6的best.pt进行YOLO推理，绘制简洁检测框后将结果图片保存到runs/cat_web_app；"
        "/api/manual允许用户不上传图片，直接根据肉眼观察选择瞳孔、耳朵和尾巴状态，由规则系统生成文本反馈。"
    )
    body(
        doc,
        f"网页端当前已生成{web_outputs}张检测结果图片，说明部署流程能够完成从浏览器上传、后端推理、检测框保存到前端显示的闭环。"
        "网页应用还加入了瞳孔状态一致性修正逻辑：当同一图像中同时出现竖瞳和圆瞳预测且高低置信度差距较大时，系统会按更高置信度结果统一处理；"
        "若差距不足，则提示用户结合肉眼观察和必要的兽医检查。"
    )
    table(
        doc,
        [
            ["接口/功能", "输入", "处理", "输出"],
            ["/", "浏览器访问", "返回HTML、CSS和JavaScript页面", "网页交互界面"],
            ["/api/predict", "上传jpg/png等图像", "YOLOv8推理、瞳孔一致性处理、检测框绘制、状态推理", "JSON结果和检测图URL"],
            ["/api/manual", "人工选择瞳孔、耳朵、尾巴状态", "规则推理与照护建议生成", "JSON文本反馈"],
            ["/outputs/<file>", "检测结果文件名", "读取runs/cat_web_app中的图片", "JPEG检测结果图"],
        ],
        widths=[3.0, 3.6, 5.2, 3.0],
    )

    doc.add_heading("第五章 实验结果与分析", level=1)
    doc.add_heading("5.1 实验设置", level=2)
    body(
        doc,
        "实验在Windows环境下完成，训练框架为Ultralytics YOLO，图像处理使用OpenCV、Pillow和Albumentations。"
        "训练使用单卡GPU，batch size设置为8，workers设置为0以适配Windows本地训练环境。"
        "所有训练均保存results.csv、训练曲线、PR曲线、F1曲线、混淆矩阵和验证集预测样例。"
    )
    doc.add_heading("5.2 多轮训练结果对比", level=2)
    body(
        doc,
        "本项目先后完成增强基线训练和多轮重点微调。增强基线cats_aug_yolov8s2的mAP50-95为0.2697，说明仅依靠增强数据训练仍难以解决细粒度局部状态识别问题。"
        "第一轮重点微调cats_focus_yolov8s将mAP50-95提升到0.6415。经过后续标签复核、数据扩充和小学习率微调，最新cats_focus_yolov8s6取得当前最佳mAP50-95为0.6472。"
    )
    table(doc, run_summary_rows(), widths=[2.2, 3.2, 1.7, 1.9, 1.9, 1.9, 2.1])
    figure(doc, BEST_RUN / "results.png", "图5-1 cats_focus_yolov8s6训练与验证曲线", 14.0)
    figure(doc, BEST_RUN / "BoxPR_curve.png", "图5-2 cats_focus_yolov8s6 PR曲线", 13.5)
    figure(doc, BEST_RUN / "BoxF1_curve.png", "图5-3 cats_focus_yolov8s6 F1曲线", 13.5)

    doc.add_heading("5.3 最新模型结果分析", level=2)
    body(
        doc,
        f"最新模型在第{best['epoch']}轮达到最佳，Precision为{float(best['metrics/precision(B)']):.4f}，Recall为{float(best['metrics/recall(B)']):.4f}，"
        f"mAP50为{float(best['metrics/mAP50(B)']):.4f}，mAP50-95为{float(best['metrics/mAP50-95(B)']):.4f}。"
        f"训练结束时第{last['epoch']}轮mAP50-95为{float(last['metrics/mAP50-95(B)']):.4f}，与最佳值接近，说明小学习率微调后模型较稳定。"
        f"相比早期重点微调最佳mAP50-95={float(old_best['metrics/mAP50-95(B)']):.4f}，最新结果提升约{float(best['metrics/mAP50-95(B)'])-float(old_best['metrics/mAP50-95(B)']):.4f}。"
    )
    figure(doc, BEST_RUN / "confusion_matrix_normalized.png", "图5-4 cats_focus_yolov8s6归一化混淆矩阵", 13.5)
    figure(doc, BEST_RUN / "val_batch0_pred.jpg", "图5-5 验证集预测结果示例（一）", 14.0)
    figure(doc, BEST_RUN / "val_batch1_pred.jpg", "图5-6 验证集预测结果示例（二）", 14.0)
    doc.add_heading("5.4 网页部署测试", level=2)
    body(
        doc,
        "网页端测试主要验证上传识别链路是否可用、模型权重路径是否正确、检测图是否能够保存并回显、手动输入模式是否能够生成建议。"
        f"从runs/cat_web_app目录看，系统已经保存{web_outputs}张网页端输出图片，证明本地网页服务已实际运行并完成多次推理。"
        "由于网页部署没有重新训练模型，因此模型精度仍以cats_focus_yolov8s6的验证集指标为准，本次更新未产生新的训练精度变化。"
    )
    if web_latest is not None:
        figure(doc, web_latest, "图5-7 网页端检测输出示例", 13.2)
    doc.add_heading("5.5 误差来源分析", level=2)
    body(
        doc,
        "从疑似错标统计和混淆矩阵可以看出，耳朵前倾与耳朵双边之间仍是主要混淆来源，圆瞳与竖瞳也存在一定类别冲突。"
        "这些问题一方面来自猫局部状态本身边界模糊，另一方面来自图像分辨率、遮挡和标注标准不完全一致。"
        "尾巴类状态受姿态和可见区域影响明显，部分情况下尾巴下垂、尾巴蜷缩和其他身体部位边界较难区分。"
    )
    body(
        doc,
        "最新实验表明，单纯增加训练轮数不一定带来持续提升；当验证指标进入平台期后，继续提升效果更依赖数据质量、类别平衡和困难样本复核。"
        "因此，标签清洗和困难样本挖掘应作为本任务的重要工程环节。"
    )

    doc.add_heading("第六章 总结与展望", level=1)
    doc.add_heading("6.1 研究总结", level=2)
    body(
        doc,
        "本文完成了基于YOLOv8的猫行为识别系统设计与实现。系统围绕猫耳、猫眼和猫尾七类局部状态建立检测流程，"
        "实现了数据集合并、离线增强、重点微调集构建、YOLOv8s训练、疑似错标发现、LabelImg人工复核同步、桌面图形界面推理和网页端应用部署。"
        "相比早期初稿，本次更新加入了数据清洗闭环、多轮训练对比和Web交互部署，使论文实验依据和应用完整性更加充分。"
    )
    body(
        doc,
        f"实验结果显示，最新模型cats_focus_yolov8s6取得mAP50={float(best['metrics/mAP50(B)']):.4f}、mAP50-95={float(best['metrics/mAP50-95(B)']):.4f}，"
        "明显优于增强基线模型。该结果说明，在猫局部状态识别这种细粒度任务中，预训练迁移、类别重采样、重点样本微调和标签质量控制具有实际效果。"
    )
    doc.add_heading("6.2 不足与展望", level=2)
    item(doc, "（1）数据规模仍需扩大：", "后续应继续采集更多品种、光照、遮挡和多猫场景图像，提高泛化能力。")
    item(doc, "（2）标注规范需进一步细化：", "耳朵和瞳孔状态存在主观边界，应制定更细的标注规则并进行多人复核。")
    item(doc, "（3）行为推理仍较简单：", "当前规则主要基于单张图片，未来可加入视频时序、姿态估计和概率推理。")
    item(doc, "（4）部署形态可扩展：", "后续可将桌面界面改造为Web端或移动端，提升实际使用便利性。")

    doc.add_heading("参考文献", level=1)
    refs = [
        "[1] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]. CVPR, 2016.",
        "[2] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal Speed and Accuracy of Object Detection[EB/OL]. arXiv:2004.10934, 2020.",
        "[3] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLO[CP/OL]. https://github.com/ultralytics/ultralytics, 2023.",
        "[4] Lin T Y, Goyal P, Girshick R, et al. Focal Loss for Dense Object Detection[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2020.",
        "[5] Lin T Y, Maire M, Belongie S, et al. Microsoft COCO: Common Objects in Context[C]. ECCV, 2014.",
        "[6] Everingham M, Van Gool L, Williams C K I, et al. The Pascal Visual Object Classes Challenge[J]. International Journal of Computer Vision, 2010.",
        "[7] Buslaev A, Iglovikov V I, Khvedchenya E, et al. Albumentations: Fast and Flexible Image Augmentations[J]. Information, 2020.",
        "[8] Bradski G. The OpenCV Library[J]. Dr. Dobb's Journal of Software Tools, 2000.",
        "[9] Ren S, He K, Girshick R, et al. Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks[J]. IEEE TPAMI, 2017.",
        "[10] Liu W, Anguelov D, Erhan D, et al. SSD: Single Shot MultiBox Detector[C]. ECCV, 2016.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        set_para_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(ref)
        font_run(r, "宋体", 12)

    doc.add_heading("致谢", level=1)
    body(
        doc,
        "在本课题研究和论文撰写过程中，感谢指导教师在选题、研究方法和论文规范方面给予的指导。"
        "感谢同学和朋友在数据收集、样本标注、复核和测试过程中提供帮助。由于本人经验有限，系统和论文仍存在不足，后续将继续完善数据规模、推理规则和应用体验。"
    )

    doc.core_properties.title = "基于YOLOv8的猫行为识别系统设计与实现"
    doc.core_properties.author = "陈祖烨"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
