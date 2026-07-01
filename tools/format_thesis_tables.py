from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


CAPTIONS = [
    "表 3.1 系统功能模块",
    "表 3.2 系统流程设计",
    "表 3.3 行为推理规则",
    "表 4.1 数据集类别分布",
    "表 4.2 疑似标签冲突统计",
    "表 4.3 标签复核批次统计",
    "表 4.4 训练参数对比",
    "表 4.5 网页服务接口设计",
    "表 5.1 多轮训练结果对比",
]


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_text_style(paragraph, size: float = 10.5, bold: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size, bold)


def remove_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)


def set_cell_margins(table, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def border(tag: str, val: str, size: int, color: str = "000000", space: int = 0) -> OxmlElement:
    elem = OxmlElement(f"w:{tag}")
    elem.set(qn("w:val"), val)
    elem.set(qn("w:sz"), str(size))
    elem.set(qn("w:space"), str(space))
    elem.set(qn("w:color"), color)
    return elem


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    # Three-line table: heavy top/bottom, light dashed internal grid.
    borders.append(border("top", "single", 12, "000000"))
    borders.append(border("bottom", "single", 12, "000000"))
    borders.append(border("left", "nil", 0, "FFFFFF"))
    borders.append(border("right", "nil", 0, "FFFFFF"))
    borders.append(border("insideH", "dashed", 4, "BFBFBF"))
    borders.append(border("insideV", "dashed", 4, "BFBFBF"))
    tbl_pr.append(borders)

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            old_tc_borders = tc_pr.find(qn("w:tcBorders"))
            if old_tc_borders is not None:
                tc_pr.remove(old_tc_borders)
            tc_borders = OxmlElement("w:tcBorders")
            if row_index == 0:
                tc_borders.append(border("bottom", "single", 8, "000000"))
            tc_pr.append(tc_borders)


def set_table_width(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9000")
    tbl_w.set(qn("w:type"), "dxa")


def format_table(table) -> None:
    set_table_width(table)
    set_table_borders(table)
    set_cell_margins(table)

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            remove_shading(cell)
            for paragraph in cell.paragraphs:
                set_paragraph_text_style(paragraph, size=10.5, bold=False)
                if row_index == 0:
                    # Match the provided sample: header centered, regular Song type.
                    for run in paragraph.runs:
                        set_run_font(run, size=10.5, bold=False)


def insert_caption_before_table(table, text: str) -> None:
    tbl = table._tbl
    prev = tbl.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        existing = "".join(node.text or "" for node in prev.iter() if node.tag == qn("w:t")).strip()
        if existing.startswith("表 "):
            for old in list(prev):
                prev.remove(old)
            prev.extend(list(make_caption_paragraph(text)))
            return

    tbl.addprevious(make_caption_paragraph(text))


def make_caption_paragraph(text: str) -> OxmlElement:
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "0")
    p_pr.append(ind)
    p.append(p_pr)

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "宋体")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    east_size = OxmlElement("w:szCs")
    east_size.set(qn("w:val"), "24")
    r_pr.append(east_size)
    r.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def format_captions(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("表 "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                set_run_font(run, 12, False)


def main() -> None:
    downloads = Path.home() / "Downloads"
    source_files = [
        p
        for p in downloads.glob("*YOLOv8*重构版.docx")
        if not p.name.startswith("~$") and "表格格式修正版" not in p.name
    ]
    if not source_files:
        raise FileNotFoundError("No thesis docx found in Downloads")
    source = sorted(source_files, key=lambda p: p.stat().st_mtime, reverse=True)[0]
    output = Path(r"C:\Users\34296\Desktop\ultralytics-main") / (source.stem + "-表格格式修正版.docx")

    doc = Document(source)
    if len(doc.tables) != len(CAPTIONS):
        print(f"warning: expected {len(CAPTIONS)} tables, found {len(doc.tables)}")

    for index, table in enumerate(doc.tables):
        if index < len(CAPTIONS):
            insert_caption_before_table(table, CAPTIONS[index])
        format_table(table)
    format_captions(doc)

    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
