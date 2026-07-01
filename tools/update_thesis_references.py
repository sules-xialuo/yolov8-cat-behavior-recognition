from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "陈祖烨-基于YOLOv8的猫行为识别系统论文初稿-重构版-表格格式修正版.docx"
OUTPUT = ROOT / "陈祖烨-基于YOLOv8的猫行为识别系统论文初稿-重构版-参考文献修正版.docx"


def set_run_font(run, size=10.5):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def set_para_font(paragraph, size=10.5):
    for run in paragraph.runs:
        set_run_font(run, size)


def append_citation(paragraph, citation):
    text = paragraph.text.rstrip()
    if text.endswith(citation):
        return
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text + citation)
    set_run_font(run)


def set_para_text(paragraph, text, size=10.5):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def set_reference_format(paragraph):
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_para_font(paragraph)


def main():
    doc = Document(SOURCE)

    citation_map = {
        62: "[1-2]",
        63: "[3-5]",
        70: "[6-10]",
        75: "[3,6]",
        77: "[6-8]",
        79: "[9]",
        117: "[8,10]",
        125: "[8-10]",
        149: "[11-13]",
        155: "[8-9]",
        159: "[11-12]",
    }
    for idx, cite in citation_map.items():
        append_citation(doc.paragraphs[idx], cite)

    # The web deployment has already been implemented; align the outlook sentence.
    set_para_text(
        doc.paragraphs[160],
        "（4）部署形态可继续扩展：当前已完成本地网页端部署，后续可进一步加入用户登录、历史记录、批量检测和移动端适配，提升实际使用便利性。",
    )

    refs = [
        "[1] Evangelista M C, Watanabe R, Leung V S Y, et al. Facial expressions of pain in cats: the development and validation of a Feline Grimace Scale[J]. Scientific Reports, 2019, 9: 19128.",
        "[2] Dawson L C, Cheal J, Niel L, et al. Humans can identify cats' affective states from subtle facial expressions[J]. Animal Welfare, 2019, 28(4): 519-531.",
        "[3] 赵永强, 饶元, 董世鹏, 等. 深度学习目标检测方法综述[J]. 中国图象图形学报, 2020, 25(4): 629-654.",
        "[4] Ren S, He K, Girshick R, et al. Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2017, 39(6): 1137-1149.",
        "[5] Lin T Y, Goyal P, Girshick R, et al. Focal Loss for Dense Object Detection[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2020, 42(2): 318-327.",
        "[6] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788.",
        "[7] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal Speed and Accuracy of Object Detection[EB/OL]. arXiv:2004.10934, 2020.",
        "[8] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLO[CP/OL]. https://github.com/ultralytics/ultralytics, 2023.",
        "[9] Buslaev A, Iglovikov V I, Khvedchenya E, et al. Albumentations: Fast and Flexible Image Augmentations[J]. Information, 2020, 11(2): 125.",
        "[10] Bradski G. The OpenCV Library[J]. Dr. Dobb's Journal of Software Tools, 2000, 25(11): 120-123.",
        "[11] Mathis A, Mamidanna P, Cury K M, et al. DeepLabCut: markerless pose estimation of user-defined body parts with deep learning[J]. Nature Neuroscience, 2018, 21(9): 1281-1289.",
        "[12] Pereira T D, Tabris N, Matsliah A, et al. SLEAP: A deep learning system for multi-animal pose tracking[J]. Nature Methods, 2022, 19(4): 486-495.",
        "[13] Nasirahmadi A, Edwards S A, Sturm B. Implementation of machine vision for detecting behaviour of cattle and pigs[J]. Livestock Science, 2017, 202: 25-38.",
    ]

    heading_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")
    ack_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "致谢")
    for p in list(doc.paragraphs[heading_idx + 1 : ack_idx]):
        remove_paragraph(p)

    heading = next(p for p in doc.paragraphs if p.text.strip() == "参考文献")
    anchor = heading
    for ref in refs:
        anchor = insert_paragraph_after(anchor, ref, style="Normal")
        set_reference_format(anchor)

    # Keep the updated file independent of the previous table-format revision.
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
