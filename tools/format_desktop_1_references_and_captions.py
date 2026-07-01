from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\34296\Desktop\1.docx")
OUTPUT = ROOT / "1-参考文献与图表题注修正版.docx"


def set_run_font(run, size=10.5):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def set_para_text(paragraph, text, size=10.5):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def set_snap_to_grid(paragraph, enabled=False):
    p_pr = paragraph._p.get_or_add_pPr()
    snap = p_pr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        p_pr.append(snap)
    snap.set(qn("w:val"), "1" if enabled else "0")


def format_reference(paragraph):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_snap_to_grid(paragraph, False)
    for run in paragraph.runs:
        set_run_font(run, 10.5)


def format_ref_heading(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_snap_to_grid(paragraph, False)


def format_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        set_run_font(run, 10.5)


def main():
    doc = Document(SOURCE)

    # Normalize obvious figure caption numbering to match the table-caption style.
    replacements = {
        "图4-1 最新重点数据集标签分布与标注框示例": "图 4.1 最新重点数据集标签分布与标注框示例",
        "图5-1 cats_focus_yolov8s6训练与验证曲线": "图 5.1 cats_focus_yolov8s6训练与验证曲线",
        "图5-2 cats_focus_yolov8s6 PR曲线": "图 5.2 cats_focus_yolov8s6 PR曲线",
        "图5-3 cats_focus_yolov8s6 F1曲线": "图 5.3 cats_focus_yolov8s6 F1曲线",
        "图5-4 cats_focus_yolov8s6归一化混淆矩阵": "图 5.4 cats_focus_yolov8s6归一化混淆矩阵",
        "图5-5 验证集预测结果示例（一）": "图 5.5 验证集预测结果示例（一）",
        "图5-6 验证集预测结果示例（二）": "图 5.6 验证集预测结果示例（二）",
        "图5-7 网页端检测输出示例": "图 5.7 网页端检测输出示例",
    }
    seen_captions = set()
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text in replacements:
            set_para_text(paragraph, replacements[text])
            text = replacements[text]
        if re.match(r"^(表|图)\s*\d+[.-]\d+", text):
            normalized_key = re.sub(r"\s+", "", text)
            if normalized_key in seen_captions:
                remove_paragraph(paragraph)
                continue
            seen_captions.add(normalized_key)
            format_caption(paragraph)
        elif text.startswith("Caption"):
            remove_paragraph(paragraph)

    refs = [
        "[1] 李林葳, 宋鑫悦, 张智盛, 等. 基于图像处理技术的动物行为识别研究进展[J]. 中国畜牧杂志, 2024, 60(10): 24-34.",
        "[2] 朱芷芫, 王海峰, 李斌, 等. 深度学习在畜禽典型行为识别中的研究进展[J]. 中国农业科技导报, 2024, 26(10): 110-124.",
        "[3] Evangelista M C, Watanabe R, Leung V S Y, et al. Facial expressions of pain in cats: the development and validation of a Feline Grimace Scale[J]. Scientific Reports, 2019, 9: 19128.",
        "[4] 赵永强, 饶元, 董世鹏, 等. 深度学习目标检测方法综述[J]. 中国图象图形学报, 2020, 25(4): 629-654.",
        "[5] 曹家乐, 李亚利, 孙汉卿, 等. 基于深度学习的视觉目标检测技术综述[J]. 中国图象图形学报, 2022, 27(6): 1697-1722.",
        "[6] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788.",
        "[7] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal Speed and Accuracy of Object Detection[EB/OL]. arXiv:2004.10934, 2020.",
        "[8] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLO[CP/OL]. https://github.com/ultralytics/ultralytics, 2023.",
        "[9] 冯晓硕, 沈樾, 王冬琦. 基于图像的数据增强方法发展现状综述[J]. 计算机科学与应用, 2021, 11(2): 370-382.",
        "[10] 郭文明, 郭鸣, 王俊, 等. 类激活映射指导数据增强的细粒度图像分类[J]. 计算机辅助设计与图形学学报, 2021, 33(11): 1698-1704.",
        "[11] Bradski G. The OpenCV Library[J]. Dr. Dobb's Journal of Software Tools, 2000, 25(11): 120-123.",
        "[12] Pereira T D, Tabris N, Matsliah A, et al. SLEAP: A deep learning system for multi-animal pose tracking[J]. Nature Methods, 2022, 19(4): 486-495.",
    ]

    heading_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")
    ack_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "致谢")
    for p in list(doc.paragraphs[heading_idx + 1 : ack_idx]):
        remove_paragraph(p)

    heading = next(p for p in doc.paragraphs if p.text.strip() == "参考文献")
    format_ref_heading(heading)
    anchor = heading
    for ref in refs:
        anchor = insert_paragraph_after(anchor, ref, style="Normal")
        format_reference(anchor)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
