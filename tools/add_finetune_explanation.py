from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "1-引用位置与参考文献格式修正版.docx"
OUTPUT = ROOT / "1-引用位置与微调说明修正版.docx"


def set_run_font(run, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def format_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        set_run_font(run)


def main():
    doc = Document(SOURCE)
    marker = "本项目先后完成增强基线训练和多轮重点微调。"
    target = next(p for p in doc.paragraphs if p.text.strip().startswith(marker))

    explanation = (
        "重点微调主要体现在训练数据、初始权重和超参数三个方面。训练数据由增强基线阶段的 "
        "yolo-cats-aug.yaml 调整为重点微调阶段的 yolo-cats-focus.yaml，该数据集加入人工挑选样本、"
        "类别重采样样本和标签复核后的样本，使模型更多关注耳朵、瞳孔和尾巴等易混淆局部目标。"
        "模型初始化不再直接使用通用 YOLOv8s 预训练权重，而是加载前一阶段较优权重 "
        "weights/cats_focus_best.pt 继续训练。在参数设置上，重点微调将学习率从基线训练的 0.001 "
        "降低到 0.00005 左右，优化器改为 AdamW，关闭 mosaic 和 mixup，减小旋转、平移、缩放等增强幅度，"
        "并设置较短的 patience，以避免在小样本细粒度任务中过度拟合训练集。"
    )
    inserted = insert_after(target, explanation, style="Normal")
    format_body(inserted)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
